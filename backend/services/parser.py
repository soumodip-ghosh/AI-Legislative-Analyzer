import io
import logging
from typing import BinaryIO

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def validate_file(filename: str, content: bytes) -> tuple[bool, str]:
    """Quick sanity checks before we do any real work."""
    if not filename:
        return False, "No filename provided."

    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        return False, f"Only PDF and DOCX files are supported. Got: '{ext}'"

    if len(content) > MAX_FILE_SIZE:
        size_mb = len(content) / (1024 * 1024)
        return False, f"File too large ({size_mb:.1f}MB). Max is 5MB."

    if len(content) < 100:
        return False, "File appears to be empty or corrupt."

    return True, ""


def extract_text_from_pdf(content: bytes) -> str:
    """Pull text from PDF pages. Skips image-only pages gracefully."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=content, filetype="pdf")
        pages = []
        for i, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages.append(f"[Page {i+1}]\n{text.strip()}")

        doc.close()

        if not pages:
            raise ValueError("No extractable text found. PDF may be image-based (scanned).")

        return "\n\n".join(pages)

    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf")


def extract_text_from_docx(content: bytes) -> str:
    """Extract paragraphs and table content from DOCX."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        chunks = []

        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve heading structure — useful for chunking later
                style = para.style.name if para.style else ""
                prefix = "## " if "Heading" in style else ""
                chunks.append(f"{prefix}{para.text.strip()}")

        # Also pull text from tables — often contains key legislative data
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    chunks.append(row_text)

        if not chunks:
            raise ValueError("No readable text found in DOCX file.")

        return "\n\n".join(chunks)

    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")


def parse_document(filename: str, content: bytes) -> str:
    """Entry point — routes to the right parser based on extension."""
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        text = extract_text_from_pdf(content)
    elif ext == "docx":
        text = extract_text_from_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    logger.info(f"Parsed '{filename}' → {len(text)} characters")
    return text
